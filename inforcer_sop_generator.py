"""
inforcer Tenant SOP Generator (Two-Phase, Selection-Driven)
===========================================================

Two-phase workflow:

  Phase 1 - EXPORT
    Pull every policy from a tenant and write an Excel selection workbook
    (<tenant>_Policy_Selection.xlsx). Each policy row has an "Include"
    tick box. Tick the policies you want in the SOP and save.

    Tick boxes use Excel 365's native checkbox cell where available and
    fall back to a plain TRUE/FALSE dropdown on older Excel.

  Phase 2 - BUILD
    Read the ticked workbook, re-fetch the tenant's current policies,
    and build the SOP Word document from the template:
      * The 7 Part 2 tables (Entra, Defender, Intune, SharePoint, Teams,
        M365 Admin Centre, Purview) keep their headers but their canonical
        rows are cleared. Whatever you ticked in the workbook gets added
        as the rows in those tables, routed by product.
      * Appendix A lists every un-ticked policy, grouped by product, so
        nothing from the tenant is silently dropped.
      * Appendix B carries the tenant metadata.

Usage
-----
    set INFORCER_API_KEY=your-key-here     (Windows - setx to persist)
    export INFORCER_API_KEY=your-key-here  (macOS/Linux)

    # Default: interactive menu (Export or Build)
    python inforcer_sop_generator.py

    # Phase 1 - export a selection workbook
    python inforcer_sop_generator.py --export-selection
    python inforcer_sop_generator.py --export-selection --tenant "Contoso"
    python inforcer_sop_generator.py --export-selection --all

    # Phase 2 - build SOP from a ticked workbook
    python inforcer_sop_generator.py --from-selection output/Contoso_Policy_Selection.xlsx

    # One-shot: include every policy without the selection step
    python inforcer_sop_generator.py --auto-match --tenant "Contoso"

Requires: Python 3.9+, `requests`, `python-docx`, `openpyxl`.
The template file must be present - defaults to `./MSP-M365-Baseline-SOP.docx`.
"""

from __future__ import annotations

import argparse
import copy
import json
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Any

import requests
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.datavalidation import DataValidation
    _HAVE_OPENPYXL = True
except ImportError:
    _HAVE_OPENPYXL = False


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

REGION_BASE_URLS = {
    "anz": "https://api-anz.inforcer.com/api",
    "eu":  "https://api-eu.inforcer.com/api",
    "uk":  "https://api-uk.inforcer.com/api",
    "us":  "https://api-us.inforcer.com/api",
}

DEFAULT_REGION = "anz"
API_KEY_ENV_VAR = "INFORCER_API_KEY"
REQUEST_TIMEOUT = 30
RETRY_BACKOFF = [2, 5, 15]

DEFAULT_TEMPLATE_PATH = "MSP-M365-Baseline-SOP.docx"

# Excel workbook layout constants
_XLSX_META_ROWS = 3            # metadata rows before the data table
_XLSX_HEADER_ROW = 5           # 1-based; row with column headers
_XLSX_FIRST_DATA_ROW = 6       # 1-based; first policy row
_XLSX_COLS = [
    "Include", "Policy Name", "Product", "Primary Group", "Secondary Group",
    "Suggested Section", "Settings Summary", "Policy ID",
]


# ---------------------------------------------------------------------------
# API client
# ---------------------------------------------------------------------------

class InforcerClient:
    """Minimal read-only client for the inforcer REST API Beta."""

    def __init__(self, api_key: str, base_url: str) -> None:
        self.base_url = base_url.rstrip("/")
        self.session = requests.Session()
        self.session.headers.update({
            "Inf-Api-Key": api_key,
            "Ocp-Apim-Subscription-Key": api_key,
            "Accept": "application/json",
        })

    def get(self, path: str, params: dict | None = None,
            soft_fail: bool = False) -> Any:
        url = f"{self.base_url}{path}"
        last_exc: Exception | None = None
        for attempt, wait in enumerate([0, *RETRY_BACKOFF]):
            if wait:
                time.sleep(wait)
            try:
                resp = self.session.get(url, params=params, timeout=REQUEST_TIMEOUT)
            except requests.RequestException as exc:
                last_exc = exc
                continue
            if resp.status_code == 200:
                payload = resp.json()
                if isinstance(payload, dict) and "data" in payload:
                    return payload["data"]
                return payload
            if resp.status_code in (429, 500, 502, 503, 504):
                continue
            if soft_fail and 400 <= resp.status_code < 500:
                return None
            raise RuntimeError(
                f"GET {url} failed: HTTP {resp.status_code} - {resp.text[:300]}"
            )
        raise RuntimeError(f"GET {url} failed after retries: {last_exc}")

    def list_tenants(self) -> list[dict]:
        data = self.get("/beta/tenants", soft_fail=True)
        return self._unwrap_list(data, ("items", "tenants", "results"))

    def get_tenant(self, tenant_id: Any) -> dict | None:
        data = self.get(f"/beta/tenants/{tenant_id}", soft_fail=True)
        return data if isinstance(data, dict) else None

    def list_tenant_policies(self, tenant_id: Any) -> list[dict]:
        data = self.get(f"/beta/tenants/{tenant_id}/policies")
        return self._unwrap_list(data, ("items", "policies", "results"))

    @staticmethod
    def _unwrap_list(data: Any, keys: tuple[str, ...]) -> list[dict]:
        if isinstance(data, list):
            return [x for x in data if isinstance(x, dict)]
        if isinstance(data, dict):
            for key in keys:
                v = data.get(key)
                if isinstance(v, list):
                    return [x for x in v if isinstance(x, dict)]
        return []


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------

def _slugify(name: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9._-]+", "_", (name or "").strip())
    return slug.strip("._-") or "tenant"


def _coerce_str(value: Any, default: str = "") -> str:
    if value is None:
        return default
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, str):
        return value if value.strip() else default
    if isinstance(value, (list, tuple)):
        return ", ".join(_coerce_str(v) for v in value) if value else default
    if isinstance(value, dict):
        for k in ("name", "displayName", "title", "id"):
            if k in value and value[k]:
                return str(value[k])
        return str(value)
    return str(value)


def _first(d: dict, *keys: str, default: Any = None) -> Any:
    for k in keys:
        v = d.get(k)
        if v not in (None, ""):
            return v
    return default


def _tenant_name(t: dict) -> str:
    return _coerce_str(_first(
        t, "tenantFriendlyName", "friendlyName", "displayName",
        "name", "tenantName", default="Unnamed tenant"
    ), default="Unnamed tenant")


def _tenant_id(t: dict) -> str:
    return _coerce_str(_first(
        t, "id", "clientTenantId", "tenantId", "msTenantId", default=""
    ))


def _tenant_domain(t: dict) -> str:
    return _coerce_str(_first(
        t, "tenantDnsName", "domain", "primaryDomain",
        "initialDomain", default=""
    ))


_POLICY_NAME_KEYS = ("displayName", "friendlyName", "name", "policyName",
                     "title", "policyDisplayName", "policyTitle")


def _deep_find(d: Any, keys: tuple[str, ...], _depth: int = 0) -> Any:
    if _depth > 4 or not isinstance(d, dict):
        return None
    for k in keys:
        if k in d and d[k] not in (None, ""):
            return d[k]
    for v in d.values():
        if isinstance(v, dict):
            r = _deep_find(v, keys, _depth + 1)
            if r not in (None, ""):
                return r
        elif isinstance(v, list):
            for item in v:
                if isinstance(item, dict):
                    r = _deep_find(item, keys, _depth + 1)
                    if r not in (None, ""):
                        return r
    return None


def _policy_name(policy: dict) -> str:
    val = _deep_find(policy, _POLICY_NAME_KEYS)
    if val:
        return str(val)
    pid = _deep_find(policy, ("id",))
    return _coerce_str(pid, default="(unnamed policy)")


def _policy_id(policy: dict) -> str:
    pid = _first(policy, "id", "policyId", "uuid")
    if pid in (None, ""):
        pid = _deep_find(policy, ("id",))
    return _coerce_str(pid)


def _policy_triple(policy: dict) -> tuple[str, str, str]:
    product = _deep_find(policy, (
        "policyCategoryProduct", "product", "productName", "productArea",
    )) or ""
    primary = _deep_find(policy, (
        "policyCategoryPrimaryGroup", "primaryGroup", "category",
    )) or ""
    secondary = _deep_find(policy, (
        "policyCategorySecondaryGroup", "secondaryGroup", "subCategory",
        "policyType",
    )) or ""
    return (str(product), str(primary), str(secondary))


def _policy_settings_summary(policy: dict, max_chars: int = 600) -> str:
    """Compact, human-readable summary of a policy's configuration."""
    skip_keys = {
        "id", "displayName", "friendlyName", "name", "policyName", "title",
        "policyDisplayName", "policyTitle",
        "policyCategoryProduct", "policyCategoryPrimaryGroup",
        "policyCategorySecondaryGroup", "product", "productName",
        "productArea", "primaryGroup", "category", "secondaryGroup",
        "subCategory",
    }
    blob: Any = None
    for k in ("policyData", "settings", "configuration", "config",
              "policySettings", "properties"):
        v = policy.get(k) if isinstance(policy, dict) else None
        if v not in (None, "", {}, []):
            blob = v
            break
    if blob is None:
        blob = {k: v for k, v in policy.items()
                if k not in skip_keys and v not in (None, "", {}, [])}
    try:
        text = json.dumps(blob, indent=2, default=str, sort_keys=True)
    except (TypeError, ValueError):
        text = str(blob)
    text = text.strip()
    if len(text) > max_chars:
        text = text[: max_chars - 1].rstrip() + "..."
    return text


def _policy_severity(policy: dict) -> str:
    """Best-effort severity extraction. Returns '' if none found."""
    val = _deep_find(policy, ("severity", "riskLevel", "impact", "priority"))
    return _coerce_str(val)


# ---------------------------------------------------------------------------
# Lay-person Control / Description helpers (used when a selected policy is
# appended to Part 2 as a row). Goal: a short, readable Control name (<= ~10
# words) and a one-sentence plain-English Description, instead of raw JSON.
# ---------------------------------------------------------------------------

# Tokens that show up in MSP policy prefixes and add no meaning for readers.
_NAME_PREFIX_NOISE = {
    "e8", "e8ibp", "ibp", "dfb", "dfc", "dfe", "dfi",
    "msp", "bp", "baseline", "baselines", "policy", "policies",
    "config", "configuration", "settings",
}

# Ordered topic detectors: (regex on lowercased name, human phrase).
# First match wins, so put more specific patterns before generic ones.
_TOPIC_PATTERNS: list[tuple[re.Pattern, str]] = [
    (re.compile(r"\bbitlocker\b", re.I),                 "BitLocker disk encryption"),
    (re.compile(r"\bfilevault\b", re.I),                 "FileVault disk encryption"),
    (re.compile(r"\bsafe\s*links?\b", re.I),             "Safe Links URL protection"),
    (re.compile(r"\bsafe\s*attachments?\b", re.I),       "Safe Attachments scanning"),
    (re.compile(r"\banti[\s-]*phish", re.I),             "anti-phishing"),
    (re.compile(r"\banti[\s-]*spam\b|\bspam\s*filter", re.I), "anti-spam filtering"),
    (re.compile(r"\banti[\s-]*malware\b|\bmalware\b", re.I),  "anti-malware"),
    (re.compile(r"\bantivirus\b|\bav\b", re.I),          "antivirus"),
    (re.compile(r"attack\s*surface|\basr\b", re.I),      "attack surface reduction"),
    (re.compile(r"\bsmartscreen\b", re.I),               "SmartScreen web/app reputation"),
    (re.compile(r"\btamper\s*protection\b", re.I),       "tamper protection"),
    (re.compile(r"\bfirewall\b", re.I),                  "firewall"),
    (re.compile(r"\bcompliance\b|\bcompliant\b", re.I),  "device compliance"),
    (re.compile(r"\bconditional\s*access\b|\bca\b", re.I), "Conditional Access"),
    (re.compile(r"\bmfa\b|multi[\s-]*factor", re.I),     "multi-factor authentication"),
    (re.compile(r"\bpassword\b", re.I),                  "password"),
    (re.compile(r"\bdlp\b|data\s*loss", re.I),           "data loss prevention"),
    (re.compile(r"\bretention\b", re.I),                 "retention"),
    (re.compile(r"sensitivity|\blabel\b", re.I),         "sensitivity labelling"),
    (re.compile(r"\bsharing\b|\bexternal\b", re.I),      "external sharing"),
    (re.compile(r"\bupdate\s*ring|\bwindows\s*update|\bwufb\b", re.I), "update ring"),
    (re.compile(r"\bautopilot\b", re.I),                 "Autopilot provisioning"),
    (re.compile(r"\bapp\s*protection\b|\bmam\b", re.I),  "app protection"),
    (re.compile(r"\bguest\b", re.I),                     "guest access"),
    (re.compile(r"\blegacy\s*auth", re.I),               "legacy authentication"),
    (re.compile(r"\bmeeting\b", re.I),                   "meeting"),
    (re.compile(r"\bmessag", re.I),                      "messaging"),
    (re.compile(r"\bquarantine\b", re.I),                "quarantine"),
    (re.compile(r"\baudit\b", re.I),                     "audit logging"),
    (re.compile(r"\bauto[\s-]*forward|remote\s*domain", re.I), "auto-forward"),
    (re.compile(r"\bdkim\b", re.I),                      "DKIM signing"),
    (re.compile(r"\bdmarc\b", re.I),                     "DMARC"),
    (re.compile(r"\blaps\b", re.I),                      "local admin password (LAPS)"),
]

# Platform detectors -> display form.
_PLATFORM_PATTERNS: list[tuple[re.Pattern, str]] = [
    (re.compile(r"\bmac[\s-]*os\b|\bmacos\b", re.I), "macOS"),
    (re.compile(r"\bwindows\b|\bwin10\b|\bwin11\b", re.I), "Windows"),
    (re.compile(r"\bios\b|\biphone\b|\bipad\b", re.I), "iOS"),
    (re.compile(r"\bandroid\b", re.I), "Android"),
    (re.compile(r"\blinux\b", re.I), "Linux"),
    (re.compile(r"\bteams\b", re.I), "Teams"),
    (re.compile(r"\bsharepoint\b|\bonedrive\b|\bspo\b|\bodb\b", re.I), "SharePoint/OneDrive"),
    (re.compile(r"\bexchange\b|\bmail\b|\bemail\b", re.I), "Exchange"),
    (re.compile(r"\bentra\b|\bazure\s*ad\b|\baad\b", re.I), "Entra ID"),
    (re.compile(r"\bintune\b", re.I), "Intune"),
    (re.compile(r"\bdefender\b", re.I), "Defender"),
]


def _detect_topic(text: str) -> str:
    for rx, phrase in _TOPIC_PATTERNS:
        if rx.search(text):
            return phrase
    return ""


def _detect_platform(text: str) -> str:
    for rx, phrase in _PLATFORM_PATTERNS:
        if rx.search(text):
            return phrase
    return ""


def _layperson_control_name(policy_name: str, max_words: int = 10) -> str:
    """Turn a raw policy name like 'E8-IBP - DfB - MacOS - Antivirus' into
    a short human label like 'macOS Antivirus policy'."""
    if not policy_name:
        return "Tenant policy"

    tokens = re.split(r"\s*[-|/]\s*|_+", policy_name)
    tokens = [t.strip() for t in tokens if t and t.strip()]

    trimmed: list[str] = []
    dropped_leading = True
    for t in tokens:
        if dropped_leading and t.lower() in _NAME_PREFIX_NOISE:
            continue
        dropped_leading = False
        trimmed.append(t)
    if not trimmed:
        trimmed = tokens[-3:] or [policy_name]

    platform = _detect_platform(policy_name)
    topic = _detect_topic(policy_name)

    if platform and topic:
        label = f"{platform} {topic} policy"
    elif topic:
        label = f"{topic.capitalize()} policy"
    elif platform:
        tail = " ".join(trimmed[-2:])
        label = f"{platform} {tail} policy" if tail else f"{platform} policy"
    else:
        tail = " ".join(trimmed[-3:])
        label = tail if "polic" in tail.lower() else f"{tail} policy"

    words = label.split()
    if len(words) > max_words:
        words = words[:max_words]
    return " ".join(words).strip() or "Tenant policy"


def _extract_settings_blob(policy: dict) -> Any:
    for k in ("policyData", "settings", "configuration", "config",
              "policySettings", "properties"):
        v = policy.get(k) if isinstance(policy, dict) else None
        if v not in (None, "", {}, []):
            return v
    return None


def _count_settings(blob: Any) -> int:
    if isinstance(blob, dict):
        return sum(1 for v in blob.values() if v not in (None, "", {}, []))
    if isinstance(blob, list):
        return len(blob)
    return 0


def _layperson_description(policy: dict, max_chars: int = 240) -> str:
    """One-sentence plain-English description of what the policy configures."""
    desc = _deep_find(policy, ("description", "policyDescription",
                               "summary", "notes"))
    if isinstance(desc, str):
        text = desc.strip()
        if text and len(text) <= max_chars:
            return text
        if text:
            return text[: max_chars - 1].rstrip() + "..."

    name = _policy_name(policy)
    product, primary, secondary = _policy_triple(policy)
    haystack = " ".join([name, product, primary, secondary])

    platform = _detect_platform(haystack)
    topic = _detect_topic(haystack)
    blob = _extract_settings_blob(policy)
    count = _count_settings(blob)

    if platform and topic:
        core = f"Configures {platform} {topic}"
    elif topic:
        core = f"Configures {topic}"
    elif platform:
        core = f"Configures {platform} policy"
    elif product:
        core = f"Configures {product} policy"
    else:
        core = "Configures tenant policy"

    if count:
        core += f" across {count} setting{'s' if count != 1 else ''}"
    core += " applied to this tenant."

    return core[: max_chars]


# ---------------------------------------------------------------------------
# Product -> Part 2 section routing
# ---------------------------------------------------------------------------

SECTION_LABELS: list[str] = [
    "2.1 Entra ID",
    "2.2 Defender for Office 365",
    "2.3 Intune",
    "2.4 SharePoint",
    "2.5 Teams",
    "2.6 M365 Admin Centre",
    "2.7 Purview",
]

SECTION_PRODUCT_PATTERNS: list[tuple[int, re.Pattern]] = [
    (0, re.compile(r"(\bentra\b|\bazure\s*ad\b|\baad\b|\bidentity\b|"
                   r"\bconditional\s*access\b)",
                   re.IGNORECASE)),
    # Purview runs before Intune/Defender so DLP/sensitivity/retention
    # policies don't get swept into 2.3 by a device-ish keyword.
    (6, re.compile(r"\b(purview|aip|information\s+protection|"
                   r"compliance\s+cent(re|er)|dlp|sensitivity|"
                   r"retention|insider\s*risk)\b",
                   re.IGNORECASE)),
    # Intune runs BEFORE the generic Defender match so endpoint-deployed
    # Defender policies (DfB, DfE, MDE, Defender for Business/Endpoint,
    # antivirus/firewall/ASR/BitLocker/FileVault/tamper protection/
    # SmartScreen/Autopilot/update rings/app protection/LAPS) land in
    # Section 2.3 Intune instead of 2.2 Defender for Office 365.
    (2, re.compile(r"(\b(intune|endpoint\s*manager|mdm|mem)\b|"
                   r"\bdf[be]\b|\bmde\b|"
                   r"\bdefender\s*for\s*(business|endpoint|cloud)\b|"
                   r"\b(bitlocker|filevault|asr|attack\s*surface)\b|"
                   r"\b(antivirus|firewall|tamper\s*protection|smartscreen)\b|"
                   r"\b(autopilot|compliance\s*polic|update\s*ring|"
                   r"app\s*protection|laps)\b)",
                   re.IGNORECASE)),
    (1, re.compile(r"(\bdfo\b|"
                   r"\bdefender\s*for\s*(office|o365)\b|"
                   r"\b(safe\s*links?|safe\s*attachments?|"
                   r"anti[\s-]*phish|anti[\s-]*spam|anti[\s-]*malware|"
                   r"preset\s*security|quarantine|zero[\s-]*hour|"
                   r"\bzap\b|\bdkim\b)|"
                   r"\bdefender\b)",
                   re.IGNORECASE)),
    (3, re.compile(r"\b(sharepoint|onedrive|spo|odb)\b", re.IGNORECASE)),
    (4, re.compile(r"\bteams\b", re.IGNORECASE)),
    (5, re.compile(r"(\bm365\s*admin\b|\bmicrosoft\s*365\s*admin\b|"
                   r"\badmin\s*cent(re|er)\b|\btenant\s*admin\b|"
                   r"\bself[-\s]?service\b|\btrials?\b|\bpurchases?\b|"
                   r"\blicen[sc]e\b|\blicen[sc]ing\b|\bexchange\b|"
                   r"\bmail\s*flow\b|\btransport\b)",
                   re.IGNORECASE)),
]


def _section_index_for_product(product: str, fallback_text: str = "") -> int | None:
    haystack = f"{product} {fallback_text}".strip()
    if not haystack:
        return None
    for idx, rx in SECTION_PRODUCT_PATTERNS:
        if rx.search(haystack):
            return idx
    return None


def _section_index_for_policy(policy: dict) -> int | None:
    """Route a policy to its Part 2 section by product/name."""
    product, primary, secondary = _policy_triple(policy)
    name = _policy_name(policy)
    return _section_index_for_product(product, f"{primary} {secondary} {name}")


def _section_label_for_policy(policy: dict) -> str:
    idx = _section_index_for_policy(policy)
    if idx is None:
        return "(Appendix - no section match)"
    return SECTION_LABELS[idx]


# ---------------------------------------------------------------------------
# Excel: write and read the selection workbook
# ---------------------------------------------------------------------------

def _require_openpyxl() -> None:
    if not _HAVE_OPENPYXL:
        raise SystemExit(
            "ERROR: the 'openpyxl' package is required for the selection "
            "workflow. Install it with: pip install openpyxl"
        )


def _coerce_tick(value: Any) -> bool:
    """Read the Include cell. Accepts native bools, TRUE/FALSE, Yes/No,
    1/0, and the string 'on' (case-insensitive)."""
    if value is None or value == "":
        return False
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    s = str(value).strip().lower()
    return s in {"true", "yes", "y", "1", "on", "x", "✓", "☑", "checked"}


# ---- Excel 365 native tick box injection ----------------------------------
#
# openpyxl can't write Excel 365's native checkbox cells directly. We save
# TRUE/FALSE booleans first (so the file always opens) and then re-zip the
# .xlsx to add the FeaturePropertyBag + cellMetadata parts that Excel 365
# uses to render those cells as clickable tick boxes. On any failure we
# leave the workbook as-is - it still works as a TRUE/FALSE dropdown.

_FPB_NS = "http://schemas.microsoft.com/office/spreadsheetml/2022/featurepropertybag"
_FPB_REL_TYPE = "http://schemas.microsoft.com/office/2022/11/relationships/FeaturePropertyBag"
_METADATA_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/sheetMetadata"
_FPB_CONTENT_TYPE = "application/vnd.ms-excel.featurepropertybag+xml"
_METADATA_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheetMetadata+xml"

_FPB_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    f'<FeaturePropertyBags xmlns="{_FPB_NS}">'
    '<bag type="Checkbox"/>'
    '</FeaturePropertyBags>'
)

_METADATA_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<metadata xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"'
    f' xmlns:xfpb="{_FPB_NS}">'
    '<metadataTypes count="1">'
    '<metadataType name="XLBPRBAG" minSupportedVersion="120000" copy="1"'
    ' pasteAll="1" pasteValues="1" merge="1" splitFirst="1" rowColShift="1"'
    ' clearFormats="1" clearComments="1" assign="1" coerce="1" cellMeta="1"/>'
    '</metadataTypes>'
    '<futureMetadata name="XLBPRBAG" count="1">'
    '<bk><extLst>'
    '<ext uri="{2bcaaf7d-bbb6-4c08-92e8-5db0d6cf2293}">'
    '<xfpb:fpb fpbi="0"/>'
    '</ext>'
    '</extLst></bk>'
    '</futureMetadata>'
    '<cellMetadata count="1">'
    '<bk><rc t="1" v="0"/></bk>'
    '</cellMetadata>'
    '</metadata>'
)


def _inject_excel365_checkboxes(xlsx_path: Path, sheet_part: str,
                                first_row: int, last_row: int,
                                column: int) -> None:
    """Convert TRUE/FALSE cells in `column` (rows first_row..last_row of
    `sheet_part`) into Excel 365 native checkbox cells. Best-effort: on
    any failure the file is left untouched."""
    import zipfile
    from xml.etree import ElementTree as ET

    if last_row < first_row:
        return

    col_letter = get_column_letter(column)
    target_cells = {f"{col_letter}{r}" for r in range(first_row, last_row + 1)}
    tmp_path = xlsx_path.with_suffix(xlsx_path.suffix + ".cbtmp")

    main_ns = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    ET.register_namespace("", main_ns)

    try:
        with zipfile.ZipFile(xlsx_path, "r") as zin:
            names = zin.namelist()
            parts = {n: zin.read(n) for n in names}

        if sheet_part not in parts:
            return

        sheet_xml = _patch_sheet_checkboxes(parts[sheet_part], target_cells, main_ns)
        if sheet_xml is None:
            return
        parts[sheet_part] = sheet_xml

        if "xl/metadata.xml" not in parts:
            parts["xl/metadata.xml"] = _METADATA_XML.encode("utf-8")

        parts["xl/featurePropertyBag/featurePropertyBag.xml"] = (
            _FPB_XML.encode("utf-8")
        )

        wb_rels_path = "xl/_rels/workbook.xml.rels"
        if wb_rels_path in parts:
            parts[wb_rels_path] = _patch_workbook_rels(parts[wb_rels_path], rels_ns)

        ct_path = "[Content_Types].xml"
        if ct_path in parts:
            parts[ct_path] = _patch_content_types(parts[ct_path], ct_ns)

        with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            written: set[str] = set()
            for name in names:
                if name in parts:
                    zout.writestr(name, parts[name])
                    written.add(name)
            for name, data in parts.items():
                if name not in written:
                    zout.writestr(name, data)

        os.replace(tmp_path, xlsx_path)
    except Exception:
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except Exception:
            pass


def _patch_sheet_checkboxes(xml_bytes: bytes, target_cells: set[str],
                            main_ns: str) -> bytes | None:
    from xml.etree import ElementTree as ET

    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return None

    ns = f"{{{main_ns}}}"
    sheet_data = root.find(f"{ns}sheetData")
    if sheet_data is None:
        return None

    touched = False
    for row_el in sheet_data.findall(f"{ns}row"):
        for c_el in row_el.findall(f"{ns}c"):
            ref = c_el.get("r")
            if ref not in target_cells:
                continue
            v_el = c_el.find(f"{ns}v")
            v_text = (v_el.text or "").strip().lower() if v_el is not None else ""
            if v_text in {"1", "true"}:
                truth = "1"
            elif v_text in {"0", "false", ""}:
                truth = "0"
            else:
                truth = "0"
            c_el.set("t", "b")
            c_el.set("cm", "1")
            for child in list(c_el):
                c_el.remove(child)
            new_v = ET.SubElement(c_el, f"{ns}v")
            new_v.text = truth
            touched = True

    if not touched:
        return None

    body = ET.tostring(root, encoding="UTF-8")
    if body.startswith(b"<?xml"):
        return body
    return b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + body


def _patch_workbook_rels(xml_bytes: bytes, rels_ns: str) -> bytes:
    from xml.etree import ElementTree as ET

    ET.register_namespace("", rels_ns)
    root = ET.fromstring(xml_bytes)
    ns = f"{{{rels_ns}}}"

    have_metadata = False
    have_fpb = False
    used_ids: set[str] = set()
    for rel in root.findall(f"{ns}Relationship"):
        rid = rel.get("Id") or ""
        used_ids.add(rid)
        rtype = rel.get("Type") or ""
        if rtype == _METADATA_REL_TYPE:
            have_metadata = True
        if rtype == _FPB_REL_TYPE:
            have_fpb = True

    def _new_id(prefix: str) -> str:
        i = 1
        while True:
            cand = f"{prefix}{i}"
            if cand not in used_ids:
                used_ids.add(cand)
                return cand
            i += 1

    if not have_metadata:
        ET.SubElement(root, f"{ns}Relationship", {
            "Id": _new_id("rIdMeta"),
            "Type": _METADATA_REL_TYPE,
            "Target": "metadata.xml",
        })
    if not have_fpb:
        ET.SubElement(root, f"{ns}Relationship", {
            "Id": _new_id("rIdFPB"),
            "Type": _FPB_REL_TYPE,
            "Target": "featurePropertyBag/featurePropertyBag.xml",
        })

    body = ET.tostring(root, encoding="UTF-8")
    if body.startswith(b"<?xml"):
        return body
    return b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + body


def _patch_content_types(xml_bytes: bytes, ct_ns: str) -> bytes:
    from xml.etree import ElementTree as ET

    ET.register_namespace("", ct_ns)
    root = ET.fromstring(xml_bytes)
    ns = f"{{{ct_ns}}}"

    have_metadata = False
    have_fpb = False
    for ov in root.findall(f"{ns}Override"):
        part = ov.get("PartName") or ""
        if part == "/xl/metadata.xml":
            have_metadata = True
        if part == "/xl/featurePropertyBag/featurePropertyBag.xml":
            have_fpb = True

    if not have_metadata:
        ET.SubElement(root, f"{ns}Override", {
            "PartName": "/xl/metadata.xml",
            "ContentType": _METADATA_CONTENT_TYPE,
        })
    if not have_fpb:
        ET.SubElement(root, f"{ns}Override", {
            "PartName": "/xl/featurePropertyBag/featurePropertyBag.xml",
            "ContentType": _FPB_CONTENT_TYPE,
        })

    body = ET.tostring(root, encoding="UTF-8")
    if body.startswith(b"<?xml"):
        return body
    return b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + body


_PRIMARY_FILL = PatternFill("solid", fgColor="0B2545") if _HAVE_OPENPYXL else None
_ACCENT_FILL  = PatternFill("solid", fgColor="D4A017") if _HAVE_OPENPYXL else None
_META_FILL    = PatternFill("solid", fgColor="EEF2F6") if _HAVE_OPENPYXL else None


def write_selection_workbook(tenant: dict, policies: list[dict],
                             region: str, out_path: Path) -> int:
    """Write the per-tenant selection workbook. Returns the number of
    policy rows written."""
    _require_openpyxl()

    wb = Workbook()
    ws = wb.active
    ws.title = "Policies"

    # ---- Metadata block (rows 1..3) -----------------------------------
    meta_rows = [
        ("Tenant",      _tenant_name(tenant)),
        ("Inforcer ID", _tenant_id(tenant)),
        ("Region",      region),
    ]
    for i, (k, v) in enumerate(meta_rows, start=1):
        kc = ws.cell(row=i, column=1, value=k)
        vc = ws.cell(row=i, column=2, value=v)
        kc.font = Font(bold=True, color="FFFFFF")
        kc.fill = _PRIMARY_FILL
        kc.alignment = Alignment(horizontal="right", vertical="center")
        vc.font = Font(bold=True)
        vc.fill = _META_FILL
        vc.alignment = Alignment(horizontal="left", vertical="center")

    # ---- Header row (row 5) --------------------------------------------
    for col_idx, header in enumerate(_XLSX_COLS, start=1):
        c = ws.cell(row=_XLSX_HEADER_ROW, column=col_idx, value=header)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = _PRIMARY_FILL
        c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # ---- Data rows -----------------------------------------------------
    settings_col = _XLSX_COLS.index("Settings Summary") + 1
    row = _XLSX_FIRST_DATA_ROW
    for p in policies:
        pid = _policy_id(p)
        name = _policy_name(p)
        product, primary, secondary = _policy_triple(p)
        sec_idx = _section_index_for_policy(p)
        suggested_section = SECTION_LABELS[sec_idx] if sec_idx is not None \
            else "(Appendix)"
        settings = _policy_settings_summary(p, max_chars=1500)

        values = [
            False, name, product, primary, secondary,
            suggested_section, settings, pid,
        ]
        for col_idx, v in enumerate(values, start=1):
            c = ws.cell(row=row, column=col_idx, value=v)
            c.alignment = Alignment(
                horizontal="center" if col_idx == 1 else "left",
                vertical="center" if col_idx == 1 else "top",
                wrap_text=(col_idx == settings_col),
            )
        ws.row_dimensions[row].height = 15
        row += 1

    # ---- Formatting: widths, freeze, autofilter, data validation ------
    widths = {1: 12, 2: 42, 3: 22, 4: 28, 5: 28, 6: 28, 7: 64, 8: 32}
    for col_idx, w in widths.items():
        ws.column_dimensions[get_column_letter(col_idx)].width = w

    ws.freeze_panes = ws.cell(row=_XLSX_FIRST_DATA_ROW, column=2)
    last_row = max(row - 1, _XLSX_FIRST_DATA_ROW)
    ws.auto_filter.ref = (
        f"A{_XLSX_HEADER_ROW}:{get_column_letter(len(_XLSX_COLS))}{last_row}"
    )

    dv = DataValidation(
        type="list", formula1='"TRUE,FALSE"', allow_blank=False,
        showDropDown=False,
    )
    dv.error = "Enter TRUE or FALSE"
    dv.errorTitle = "Invalid Include value"
    dv.prompt = "Tick the box to include this policy in the SOP"
    dv.promptTitle = "Include"
    ws.add_data_validation(dv)
    if last_row >= _XLSX_FIRST_DATA_ROW:
        dv.add(f"A{_XLSX_FIRST_DATA_ROW}:A{last_row}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    _save_workbook_safely(wb, out_path)

    if last_row >= _XLSX_FIRST_DATA_ROW:
        _inject_excel365_checkboxes(
            out_path,
            sheet_part="xl/worksheets/sheet1.xml",
            first_row=_XLSX_FIRST_DATA_ROW,
            last_row=last_row,
            column=1,
        )
    return max(row - _XLSX_FIRST_DATA_ROW, 0)


def _save_workbook_safely(wb: Workbook, out_path: Path,
                          retries: int = 6, delay: float = 1.5) -> None:
    tmp_path = out_path.with_suffix(out_path.suffix + ".tmp")
    last_exc: Exception | None = None
    for attempt in range(retries):
        try:
            wb.save(str(tmp_path))
            os.replace(tmp_path, out_path)
            return
        except PermissionError as exc:
            last_exc = exc
            try:
                if tmp_path.exists():
                    tmp_path.unlink()
            except Exception:
                pass
            if attempt < retries - 1:
                time.sleep(delay * (attempt + 1))
                continue
            raise
    if last_exc:
        raise last_exc


def read_selection_workbook(xlsx_path: Path) -> dict:
    """Read a selection workbook. Returns:
       {
         "tenant_name": str,
         "tenant_id":   str,
         "region":      str,
         "selected_ids": set[str],        # policies whose tick box is on
         "all_rows":     list[dict],      # every row with its fields
       }
    """
    _require_openpyxl()

    wb = load_workbook(str(xlsx_path), data_only=True, read_only=False)
    if "Policies" not in wb.sheetnames:
        raise SystemExit(
            f"ERROR: {xlsx_path} does not look like a selection workbook "
            "(missing 'Policies' sheet)."
        )
    ws = wb["Policies"]

    def _cell_text(row: int, col: int) -> str:
        v = ws.cell(row=row, column=col).value
        return str(v).strip() if v is not None else ""

    meta = {
        "tenant_name": _cell_text(1, 2),
        "tenant_id":   _cell_text(2, 2),
        "region":      _cell_text(3, 2) or DEFAULT_REGION,
    }

    expected = [c.lower() for c in _XLSX_COLS]
    actual = [_cell_text(_XLSX_HEADER_ROW, i + 1).lower()
              for i in range(len(_XLSX_COLS))]
    if actual != expected:
        raise SystemExit(
            f"ERROR: {xlsx_path} header row does not match. Expected "
            f"{_XLSX_COLS}, got {[ws.cell(row=_XLSX_HEADER_ROW, column=i+1).value for i in range(len(_XLSX_COLS))]}"
        )

    pid_col = _XLSX_COLS.index("Policy ID") + 1
    settings_col = _XLSX_COLS.index("Settings Summary") + 1

    selected_ids: set[str] = set()
    all_rows: list[dict] = []
    row = _XLSX_FIRST_DATA_ROW
    while True:
        name = _cell_text(row, 2)
        pid = _cell_text(row, pid_col)
        if not name and not pid:
            break
        include_raw = ws.cell(row=row, column=1).value
        rec = {
            "include":           _coerce_tick(include_raw),
            "name":              name,
            "product":           _cell_text(row, 3),
            "primary":           _cell_text(row, 4),
            "secondary":         _cell_text(row, 5),
            "suggested_section": _cell_text(row, 6),
            "settings_summary":  _cell_text(row, settings_col),
            "policy_id":         pid,
        }
        all_rows.append(rec)
        if rec["include"] and pid:
            selected_ids.add(pid)
        row += 1

    return {
        **meta,
        "selected_ids": selected_ids,
        "all_rows":     all_rows,
    }


# ---------------------------------------------------------------------------
# Word template manipulation
# ---------------------------------------------------------------------------

def _replace_paragraph_text(paragraph, new_text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def _update_title_block(doc: Document, tenant: dict, policy_count: int,
                        selected_count: int | None = None) -> None:
    name = _tenant_name(tenant)
    domain = _tenant_domain(tenant)
    tid = _tenant_id(tenant)
    version = datetime.now().strftime("%Y.%m")

    bits = [f"Tenant: {name}"]
    if domain:
        bits.append(f"Domain: {domain}")
    if selected_count is not None:
        bits.append(f"Controls selected: {selected_count} of {policy_count}")
    else:
        bits.append(f"Policies captured: {policy_count}")
    if tid:
        bits.append(f"Inforcer tenant ID: {tid}")
    bits.append(f"v{version}")
    stamp = "  -  ".join(bits)

    if len(doc.paragraphs) > 3:
        _replace_paragraph_text(doc.paragraphs[3], stamp)


def _shade_cell(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _apply_table_borders(table, color_hex: str = "808080", size: str = "4") -> None:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tbl_pr.append(borders)


_PART2_HEADER = ("control", "description", "severity", "policy display name")


def _find_part2_tables(doc: Document) -> list:
    """Return the Part 2 tables in document order. Matches on header row."""
    out = []
    for table in doc.tables:
        if not table.rows:
            continue
        hdr = tuple(
            re.sub(r"\s+", " ", c.text).strip().lower()
            for c in table.rows[0].cells
        )
        if hdr[:4] == _PART2_HEADER:
            out.append(table)
    return out


def _bucket_policies_by_section(policies: list[dict]) -> dict[int, list[dict]]:
    by_section: dict[int, list[dict]] = {i: [] for i in range(7)}
    by_section[-1] = []
    for p in policies:
        idx = _section_index_for_policy(p)
        if idx is None or idx < 0 or idx >= 7:
            by_section[-1].append(p)
        else:
            by_section[idx].append(p)
    return by_section


def _set_cell_text(cell, text: str) -> None:
    """Clear a cell and set its text, preserving the first paragraph's style."""
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        cell.text = text
        return
    first = paragraphs[0]
    runs = first.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)


def _populate_section_tables(doc: Document,
                             selected: list[dict]
                             ) -> set[int]:
    """Populate Part 2 tables directly from the policies the user ticked.

    For each of the 7 section tables we keep the header row, throw away
    every canonical/template row underneath it, then add one row per
    ticked policy that routes to that section.

    Returns the set of section indices that ended up with content.
    """
    section_tables = _find_part2_tables(doc)
    if len(section_tables) < 7:
        print(f"  WARN: expected 7 Part 2 tables, found {len(section_tables)} - "
              "template may be out of date. Extra sections will be skipped.")

    by_section_selected = _bucket_policies_by_section(selected)
    populated_sections: set[int] = set()

    for sec_idx, table in enumerate(section_tables):
        if sec_idx >= 7:
            break

        rows = list(table.rows)
        tbl_el = table._tbl

        # Clone the first data row's XML before deleting it, so newly
        # appended rows can inherit the canonical formatting (font,
        # paragraph style, cell widths, borders, shading) instead of
        # the bare-bones layout that table.add_row() produces.
        template_row_el = (
            copy.deepcopy(rows[1]._tr) if len(rows) >= 2 else None
        )

        # Delete every existing data row (keep the header row only).
        for row in rows[1:]:
            tbl_el.remove(row._tr)

        # Append one row per ticked policy routed to this section.
        candidates = by_section_selected.get(sec_idx, [])
        for pol in sorted(candidates, key=lambda x: _policy_name(x).lower()):
            if template_row_el is not None:
                new_tr = copy.deepcopy(template_row_el)
                tbl_el.append(new_tr)
                new_row = table.rows[-1]
            else:
                new_row = table.add_row()
            cells = new_row.cells
            if len(cells) < 4:
                continue
            _set_cell_text(cells[0], _layperson_control_name(_policy_name(pol)))
            _set_cell_text(cells[1], _layperson_description(pol))
            _set_cell_text(cells[2], _policy_severity(pol) or "-")
            _set_cell_text(cells[3], _policy_name(pol))
            populated_sections.add(sec_idx)

    return populated_sections


# ---------------------------------------------------------------------------
# Appendices
# ---------------------------------------------------------------------------

def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _append_excluded_policies_appendix(doc: Document,
                                       excluded: list[dict],
                                       tenant_name: str) -> None:
    """Appendix A - tenant policies NOT ticked for inclusion in Part 2."""
    doc.add_page_break()
    _add_heading(doc, "Appendix A - Tenant Policies Not Included in SOP", level=1)
    doc.add_paragraph(
        f"Policies present in tenant \"{tenant_name}\" that were not selected "
        "for inclusion in the SOP above. They are listed here so nothing is "
        "silently dropped - review this list when the SOP is next revised to "
        "decide whether any should be promoted into Part 2."
    )

    if not excluded:
        doc.add_paragraph(
            "Every tenant policy was included in the SOP. Nothing to list."
        )
        return

    by_product: dict[str, list[dict]] = {}
    for p in excluded:
        prod = _policy_triple(p)[0] or "Unspecified"
        by_product.setdefault(prod, []).append(p)

    for product in sorted(by_product.keys()):
        doc.add_paragraph().add_run(product).bold = True
        table = doc.add_table(rows=1, cols=3)
        _apply_table_borders(table)
        hdr = table.rows[0].cells
        for i, h in enumerate(["Policy Name", "Primary Group", "Secondary Group"]):
            hdr[i].text = h
            _shade_cell(hdr[i], "0B2545")
            for para in hdr[i].paragraphs:
                for r in para.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for p in sorted(by_product[product], key=lambda x: _policy_name(x).lower()):
            _, prim, sec = _policy_triple(p)
            row = table.add_row().cells
            row[0].text = _policy_name(p)
            row[1].text = prim or "-"
            row[2].text = sec or "-"
        doc.add_paragraph()


def _append_tenant_metadata_appendix(doc: Document, tenant: dict,
                                     total_policies: int,
                                     selected_count: int) -> None:
    doc.add_page_break()
    _add_heading(doc, "Appendix B - Tenant Metadata", level=1)

    pairs: list[tuple[str, Any]] = [
        ("Tenant Name",            _tenant_name(tenant)),
        ("Tenant Domain",          _tenant_domain(tenant)),
        ("Inforcer Tenant ID",     _tenant_id(tenant)),
        ("Microsoft Tenant ID",    _first(tenant, "msTenantId", "microsoftTenantId")),
        ("Total policies in tenant", total_policies),
        ("Controls included in SOP", selected_count),
        ("Controls in appendix",     total_policies - selected_count),
        ("Report generated",       datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]
    pairs = [(k, v) for k, v in pairs if v is not None and v != ""]

    if pairs:
        table = doc.add_table(rows=len(pairs), cols=2)
        _apply_table_borders(table)
        for i, (k, v) in enumerate(pairs):
            row = table.rows[i]
            row.cells[0].text = str(k)
            row.cells[1].text = _coerce_str(v, default="-")
            for p in row.cells[0].paragraphs:
                for r in p.runs:
                    r.bold = True
        doc.add_paragraph()


def _save_doc_safely(doc: Document, out_path: Path,
                     retries: int = 6, delay: float = 1.5) -> None:
    tmp_path = out_path.with_suffix(out_path.suffix + ".tmp")
    last_exc: Exception | None = None
    for attempt in range(retries):
        try:
            doc.save(str(tmp_path))
            os.replace(tmp_path, out_path)
            return
        except PermissionError as exc:
            last_exc = exc
            try:
                if tmp_path.exists():
                    tmp_path.unlink()
            except Exception:
                pass
            if attempt < retries - 1:
                time.sleep(delay * (attempt + 1))
                continue
            raise
    if last_exc:
        raise last_exc


# ---------------------------------------------------------------------------
# SOP building
# ---------------------------------------------------------------------------

def build_sop(template_path: Path, tenant: dict,
              selected: list[dict], excluded: list[dict],
              out_path: Path) -> None:
    """Build an SOP for a single tenant using a pre-partitioned list of
    selected (include in Part 2) and excluded (listed in Appendix A) policies."""
    doc = Document(str(template_path))
    total = len(selected) + len(excluded)
    _update_title_block(doc, tenant, total, selected_count=len(selected))
    _populate_section_tables(doc, selected)
    _append_excluded_policies_appendix(doc, excluded, _tenant_name(tenant))
    _append_tenant_metadata_appendix(doc, tenant, total, len(selected))
    _save_doc_safely(doc, out_path)


# ---------------------------------------------------------------------------
# Tenant selection (CLI + interactive picker)
# ---------------------------------------------------------------------------

def _match_tenant_filter(tenants: list[dict], needle: str) -> list[dict]:
    n = needle.strip()
    if not n:
        return []
    exact: list[dict] = []
    nl = n.lower()
    for t in tenants:
        if _tenant_id(t) == n or _tenant_name(t) == n or _tenant_domain(t) == n:
            exact.append(t)
    if exact:
        return exact
    return [t for t in tenants
            if nl in _tenant_name(t).lower()
            or nl in _tenant_domain(t).lower()]


def _resolve_cli_selection(tenants: list[dict], needles: list[str]) -> list[dict]:
    selected: list[dict] = []
    seen_ids: set[str] = set()
    unresolved: list[str] = []
    for needle in needles:
        matches = _match_tenant_filter(tenants, needle)
        if not matches:
            unresolved.append(needle)
            continue
        for m in matches:
            tid = _tenant_id(m)
            if tid in seen_ids:
                continue
            seen_ids.add(tid)
            selected.append(m)
    if unresolved:
        raise SystemExit(
            "ERROR: no tenant matched: " + ", ".join(repr(u) for u in unresolved)
        )
    return selected


def _prompt_tenant_selection(tenants: list[dict]) -> list[dict]:
    print()
    print("Available tenants:")
    for i, t in enumerate(tenants, start=1):
        name = _tenant_name(t)
        dom = _tenant_domain(t)
        tid = _tenant_id(t)
        suffix = f"({dom}, id {tid})" if dom else f"(id {tid})"
        print(f"  [{i:>2}] {name}   {suffix}")
    print("  [all] Every tenant above")
    print()

    while True:
        try:
            raw = input(
                "Which tenant(s)? (e.g. 1, 1,3,5, or all): "
            ).strip()
        except EOFError:
            raise SystemExit(
                "ERROR: no input available. Use --tenant or --all for "
                "non-interactive runs."
            )
        if not raw:
            continue
        if raw.lower() in ("all", "*"):
            return list(tenants)
        tokens = [t for t in re.split(r"[,\s]+", raw) if t]
        try:
            indices = [int(t) for t in tokens]
        except ValueError:
            print("  Please enter numbers (e.g. 1,3,5) or 'all'.")
            continue
        if any(i < 1 or i > len(tenants) for i in indices):
            print(f"  Numbers must be between 1 and {len(tenants)}.")
            continue
        seen: set[int] = set()
        selection: list[dict] = []
        for i in indices:
            if i in seen:
                continue
            seen.add(i)
            selection.append(tenants[i - 1])
        return selection


def _prompt_mode_choice() -> str:
    """Ask whether to export a selection workbook or build an SOP from one."""
    print()
    print("What do you want to do?")
    print("  [1] Export policies to an Excel selection file (Phase 1)")
    print("  [2] Build an SOP from an Excel selection file you have already "
          "reviewed (Phase 2)")
    print("  [Q] Quit")
    while True:
        try:
            raw = input("Enter 1, 2, or Q: ").strip().lower()
        except EOFError:
            raise SystemExit(
                "ERROR: no input available. Use --export-selection or "
                "--from-selection for non-interactive runs."
            )
        if raw in ("q", "quit", "exit"):
            raise SystemExit(0)
        if raw == "1":
            return "export"
        if raw == "2":
            return "build"
        print("  Please enter 1, 2, or Q.")


def _prompt_selection_file(out_dir: Path) -> Path:
    candidates = sorted(out_dir.glob("*_Policy_Selection.xlsx")) \
        if out_dir.exists() else []
    if candidates:
        print()
        print("Selection workbooks found in output folder:")
        for i, path in enumerate(candidates, start=1):
            print(f"  [{i:>2}] {path.name}")
        print("  [O] Enter another path")
        while True:
            try:
                raw = input("Which workbook? (number or O): ").strip().lower()
            except EOFError:
                raise SystemExit(
                    "ERROR: no input available. Use --from-selection PATH "
                    "for non-interactive runs."
                )
            if raw == "o":
                break
            if raw.isdigit():
                idx = int(raw)
                if 1 <= idx <= len(candidates):
                    return candidates[idx - 1]
            print("  Please enter a listed number or O.")
    while True:
        try:
            raw = input("Path to selection workbook (.xlsx): ").strip()
        except EOFError:
            raise SystemExit(
                "ERROR: no input available. Use --from-selection PATH "
                "for non-interactive runs."
            )
        if not raw:
            continue
        path = Path(raw)
        if path.exists():
            return path
        print(f"  Not found: {path}")


# ---------------------------------------------------------------------------
# Workflow phases
# ---------------------------------------------------------------------------

def run_export_phase(client: InforcerClient, region: str, out_dir: Path,
                     tenants_to_export: list[dict]) -> list[Path]:
    """Phase 1: for each selected tenant, fetch policies and write an Excel
    selection workbook. Returns the list of workbooks written."""
    _require_openpyxl()
    written: list[Path] = []
    seen_slugs: dict[str, int] = {}

    for t in tenants_to_export:
        name = _tenant_name(t)
        tid = _first(t, "id", "clientTenantId", "tenantId", default=None)
        if tid is None:
            print(f"  {name}: no tenant ID available - skipped.")
            continue
        try:
            policies = client.list_tenant_policies(tid)
        except Exception as exc:
            short = str(exc).split("-", 1)[0].strip()[:140]
            print(f"  {name}: policy fetch failed - {short}")
            continue

        slug = _slugify(name)
        count = seen_slugs.get(slug, 0)
        seen_slugs[slug] = count + 1
        fname = (f"{slug}_Policy_Selection.xlsx" if count == 0
                 else f"{slug}_Policy_Selection_{count + 1}.xlsx")
        out_path = out_dir / fname
        try:
            written_count = write_selection_workbook(t, policies, region, out_path)
            print(f"  wrote {out_path}")
            print(f"         policies written: {written_count}")
            written.append(out_path)
        except PermissionError as exc:
            print(f"  SKIPPED {out_path} - file locked ({exc}).")
        except Exception as exc:
            print(f"  SKIPPED {out_path} - {exc}")

    return written


def run_build_phase(api_key: str, xlsx_path: Path, template_path: Path,
                    out_dir: Path) -> Path | None:
    """Phase 2: read the workbook, re-fetch the tenant's policies, build
    the SOP using ticked rows for Part 2 and un-ticked rows for Appendix A."""
    sel = read_selection_workbook(xlsx_path)
    region = (sel.get("region") or DEFAULT_REGION).lower()
    if region not in REGION_BASE_URLS:
        print(f"  WARN: workbook region '{region}' is unknown, falling back to {DEFAULT_REGION}.")
        region = DEFAULT_REGION
    base_url = REGION_BASE_URLS[region]

    tenant_id = sel["tenant_id"]
    if not tenant_id:
        raise SystemExit(
            f"ERROR: workbook {xlsx_path} has no Inforcer tenant ID in the "
            "metadata block. Re-export and retry."
        )

    print(f"Using workbook: {xlsx_path}")
    print(f"  Tenant: {sel['tenant_name']} (id {tenant_id}, region {region})")
    print(f"  Policies ticked for inclusion: {len(sel['selected_ids'])} / "
          f"{len(sel['all_rows'])}")

    client = InforcerClient(api_key, base_url)

    tenant = client.get_tenant(tenant_id) or {
        "id": tenant_id,
        "displayName": sel["tenant_name"],
    }
    try:
        policies = client.list_tenant_policies(tenant_id)
    except Exception as exc:
        raise SystemExit(f"ERROR: failed to re-fetch tenant policies: {exc}")

    print(f"  Re-fetched {len(policies)} policies from Inforcer.")

    selected_ids = sel["selected_ids"]
    by_id = {_policy_id(p): p for p in policies}
    selected: list[dict] = [by_id[pid] for pid in selected_ids if pid in by_id]
    excluded: list[dict] = [p for p in policies
                            if _policy_id(p) not in selected_ids]

    missing = selected_ids - set(by_id.keys())
    if missing:
        print(f"  NOTE: {len(missing)} ticked policy IDs no longer exist in the "
              "tenant and will be skipped.")

    slug = _slugify(sel["tenant_name"])
    out_path = out_dir / f"{slug}_SOP.docx"
    try:
        build_sop(template_path, tenant, selected, excluded, out_path)
    except PermissionError as exc:
        print(f"  SKIPPED {out_path} - file locked ({exc}).")
        return None
    except Exception as exc:
        print(f"  SKIPPED {out_path} - {exc}")
        return None

    print(f"  wrote {out_path}")
    print(f"         controls in SOP: {len(selected)}   "
          f"policies in appendix: {len(excluded)}")
    return out_path


def run_auto_match_phase(client: InforcerClient, template_path: Path,
                         out_dir: Path, tenants: list[dict]) -> None:
    """One-shot: fetch every policy and treat them all as selected. Skips
    the Excel review step entirely - useful for first-pass drafts."""
    for t in tenants:
        name = _tenant_name(t)
        tid = _first(t, "id", "clientTenantId", "tenantId", default=None)
        if tid is None:
            print(f"  {name}: no tenant ID available - skipped.")
            continue
        try:
            policies = client.list_tenant_policies(tid)
            print(f"  {name} (id={tid}): {len(policies)} policies")
        except Exception as exc:
            short = str(exc).split("-", 1)[0].strip()[:140]
            print(f"  {name}: policy fetch failed - {short}")
            continue
        slug = _slugify(name)
        out_path = out_dir / f"{slug}_SOP.docx"
        try:
            build_sop(template_path, t, selected=policies, excluded=[],
                      out_path=out_path)
            print(f"  wrote {out_path}")
        except Exception as exc:
            print(f"  SKIPPED {out_path} - {exc}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Generate an SOP Word document from a tenant's policies "
                    "in Inforcer using a two-phase selection workflow."
    )
    p.add_argument("--region", choices=list(REGION_BASE_URLS.keys()),
                   default=DEFAULT_REGION,
                   help=f"inforcer region (default: {DEFAULT_REGION}).")
    p.add_argument("--out", default="./output",
                   help="Output directory (default: ./output).")
    p.add_argument("--template", default=DEFAULT_TEMPLATE_PATH,
                   help=f"Path to the SOP template .docx (default: {DEFAULT_TEMPLATE_PATH}).")
    p.add_argument("--tenant", action="append", default=[],
                   help="Tenant name, ID, domain, or name substring. Repeat for multiple. "
                        "Applies to --export-selection and --auto-match.")
    p.add_argument("--all", action="store_true",
                   help="Include every tenant (skips interactive picker).")
    p.add_argument("--export-selection", action="store_true",
                   help="Phase 1: fetch tenant policies and write an Excel "
                        "selection workbook for each selected tenant.")
    p.add_argument("--from-selection", action="append", default=[],
                   help="Phase 2: build an SOP from a ticked selection workbook. "
                        "Repeat to build SOPs from multiple workbooks.")
    p.add_argument("--auto-match", action="store_true",
                   help="One-shot: build SOPs directly without the selection step "
                        "(every policy is included in Part 2).")
    return p.parse_args()


def main() -> int:
    args = parse_args()

    api_key = os.environ.get(API_KEY_ENV_VAR)
    if not api_key and not args.from_selection:
        print(f"ERROR: environment variable {API_KEY_ENV_VAR} is not set.", file=sys.stderr)
        print(f"       set {API_KEY_ENV_VAR}=your-key-here and re-run.", file=sys.stderr)
        return 2

    template_path = Path(args.template)
    if not template_path.exists():
        alt = Path(__file__).parent / DEFAULT_TEMPLATE_PATH
        if alt.exists():
            template_path = alt
        else:
            print(f"ERROR: template file not found: {args.template}", file=sys.stderr)
            print(f"       Also tried: {alt}", file=sys.stderr)
            return 2

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    base_url = REGION_BASE_URLS[args.region]

    # ---- Route to a workflow phase ------------------------------------
    mode: str
    if args.from_selection:
        mode = "build"
    elif args.export_selection:
        mode = "export"
    elif args.auto_match:
        mode = "auto"
    else:
        mode = _prompt_mode_choice()

    # ---- Phase 2: build from selection file(s) ------------------------
    if mode == "build":
        if not api_key:
            print(f"ERROR: {API_KEY_ENV_VAR} must be set to re-fetch policies.",
                  file=sys.stderr)
            return 2
        print(f"Using template: {template_path}")
        xlsx_paths: list[Path]
        if args.from_selection:
            xlsx_paths = [Path(p) for p in args.from_selection]
        else:
            xlsx_paths = [_prompt_selection_file(out_dir)]

        for xp in xlsx_paths:
            if not xp.exists():
                print(f"  SKIPPED {xp} - file not found.")
                continue
            try:
                run_build_phase(api_key, xp, template_path, out_dir)
            except SystemExit:
                raise
            except Exception as exc:
                print(f"  SKIPPED {xp} - {exc}")
        print("Done.")
        return 0

    # ---- Phase 1 and one-shot auto-match both need the tenant list ----
    client = InforcerClient(api_key, base_url)
    print(f"Fetching tenants from {base_url}/beta/tenants ...")
    try:
        tenants = client.list_tenants()
    except Exception as exc:
        print(f"ERROR: failed to fetch tenants: {exc}", file=sys.stderr)
        return 1
    if not tenants:
        print("No tenants returned. Check your API key permissions.")
        return 0
    print(f"Retrieved {len(tenants)} tenant(s).")

    if args.all:
        selected = list(tenants)
        print(f"--all specified: including every tenant ({len(selected)}).")
    elif args.tenant:
        selected = _resolve_cli_selection(tenants, args.tenant)
        print(f"--tenant matched {len(selected)} tenant(s): "
              + ", ".join(_tenant_name(t) for t in selected))
    else:
        selected = _prompt_tenant_selection(tenants)
        print(f"Selected {len(selected)} tenant(s).")

    if not selected:
        print("No tenants selected. Nothing to do.")
        return 0

    if mode == "export":
        print(f"Using template: {template_path}")
        print("Fetching policies and writing selection workbooks...")
        written = run_export_phase(client, args.region, out_dir, selected)
        print("Done.")
        if written:
            print()
            print("Next step:")
            print("  1. Open each workbook in Excel.")
            print("  2. Tick the 'Include' column for the policies you want "
                  "in the SOP.")
            print("  3. Save the workbook when you're happy.")
            print("  4. Re-run to build the SOP - e.g.:")
            for p in written:
                rel = p.as_posix()
                print(f'        python inforcer_sop_generator.py --from-selection "{rel}"')
        return 0

    # mode == "auto"
    print(f"Using template: {template_path}")
    print("Auto-match mode - including every policy in the SOP.")
    run_auto_match_phase(client, template_path, out_dir, selected)
    print("Done.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
